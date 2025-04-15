#!/usr/bin/env python3
"""
多功能文档比对工具 v2.1
支持格式：Microsoft Word (.docx), Adobe PDF (.pdf)
"""

import os
import re
import difflib
import argparse
from collections import defaultdict
from typing import List, Dict,Tuple  
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog

import webbrowser
# Word处理依赖
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import uuid
# PDF处理依赖
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.colors import yellow, cyan
import json
from difflib import SequenceMatcher
import urllib.parse
from reportlab.lib import colors
import pdfplumber
import logging
import warnings
from docx.shared import RGBColor  # 确保正确导入 RGBColor
# 忽略pdfminer生成的特定警告
warnings.filterwarnings("ignore", category=UserWarning, message="CropBox missing from /Page, defaulting to MediaBox")

# ================== 通用工具函数 ==================
def highlight_text(paragraph, start_pos, end_pos, highlight_color=WD_COLOR_INDEX.YELLOW):
    current_pos = 0
    runs = list(paragraph.runs)

    for run in runs:
        run_text = run.text
        run_length = len(run_text)
        
        if current_pos <= start_pos < current_pos + run_length:
            split_pos_start = start_pos - current_pos
            split_pos_end = end_pos - current_pos
            
            if split_pos_start > 0:
                new_run = paragraph.add_run(run_text[:split_pos_start])
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.color.rgb = run.font.color.rgb
            
            highlight_run = paragraph.add_run(run_text[split_pos_start:split_pos_end])
            highlight_run.font.highlight_color = highlight_color
            highlight_run.bold = run.bold
            highlight_run.italic = run.italic
            
            if split_pos_end < run_length:
                new_run = paragraph.add_run(run_text[split_pos_end:])
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.color.rgb = run.font.color.rgb
            
            paragraph._p.remove(run._r)
            continue
        
        elif current_pos + run_length <= start_pos:
            new_run = paragraph.add_run(run_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.font.color.rgb = run.font.color.rgb
        
        else:
            new_run = paragraph.add_run(run_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.font.color.rgb = run.font.color.rgb
        
        current_pos += run_length
        paragraph._p.remove(run._r)

# ================== Word对比模块 ==================
def compare_docs_with_threshold(doc1_path, doc2_path, min_length=15):
    """增强版文档对比函数，确保生成完整的位置信息"""
    try:
        doc1 = Document(doc1_path)
        doc2 = Document(doc2_path)
    except Exception as e:
        print(f"❌ 无法打开文档: {e}")
        return []

    matches = []
    
    # 段落级对比
    for i, para1 in enumerate(doc1.paragraphs):
        for j, para2 in enumerate(doc2.paragraphs):
            # 新增段落文本清洗
            text1 = re.sub(r'\s+', '', para1.text)
            text2 = re.sub(r'\s+', '', para2.text)
            
            if text1 == text2 and len(text1) >= min_length:
              
                matches.append({
                    "type": "full_match",
                    "doc1_para": i,
                    "doc1_pos": (0, len(para1.text)),  # 新增全段落位置
                    "doc1_page": get_page_number(doc1, i),
                    "doc1_line": get_line_number(doc1, i),
                    "doc2_para": j,
                    "doc2_pos": (0, len(para2.text)),
                    "doc2_page": get_page_number(doc2, j),
                    "doc2_line": get_line_number(doc2, j)
                })
            else:
                # 使用改进的LCS算法进行局部匹配
                seq_matcher = difflib.SequenceMatcher(None, text1, text2)
                for match in seq_matcher.get_matching_blocks():
                    if match.size >= min_length:
                        matches.append({
                            "type": "partial_match",
                            "doc1_para": i,
                            "doc1_pos": (match.a, match.a + match.size),
                            "doc1_page": get_page_number(doc1, i),
                            "doc1_line": get_line_number(doc1, i),
                            "doc2_para": j,
                            "doc2_pos": (match.b, match.b + match.size),
                            "doc2_page": get_page_number(doc2, j),
                            "doc2_line": get_line_number(doc2, j)
                        })
    return matches

def get_page_number(doc, para_index):
    """获取段落在文档中的页码"""
    # 这里假设每个段落都在新的一页，实际情况可能需要更复杂的逻辑
    return para_index // 50 + 1  # 假设每页最多50个段落

def get_line_number(doc, para_index):
    """获取段落在文档中的行号"""
    # 这里假设每个段落都在新的一行，实际情况可能需要更复杂的逻辑
    return para_index % 50 + 1  # 假设每页最多50个段落

def add_comment(paragraph, comment_text):
    """在段落后添加注释"""
    # 获取段落的 XML
    p = paragraph._element
    # 创建注释元素
    comment_id = str(uuid.uuid4())  # 注释 ID，使用 UUID 生成唯一 ID
    comment = etree.Element(qn('w:comment'), {
        qn('w:id'): comment_id,
        qn('w:author'): '匹配工具',
        qn('w:date'): '2023-10-01T00:00:00Z',  # 日期时间，实际使用时需要动态生成
    })
    comment_p = etree.SubElement(comment, qn('w:p'))
    comment_run = etree.SubElement(comment_p, qn('w:r'))
    comment_text_elem = etree.SubElement(comment_run, qn('w:t'))
    comment_text_elem.text = comment_text

    # 将注释添加到文档的注释部分
    comments_part = paragraph.part
    if not hasattr(comments_part, '_comments'):
        comments_part._comments = etree.Element(qn('w:comments'))
    comments_part._comments.append(comment)

    # 创建引用注释的元素
    comment_reference = etree.Element(qn('w:commentReference'), {
        qn('w:id'): comment_id,
    })
    p.append(comment_reference)


def add_bookmark(paragraph, bookmark_name):
    """在段落后添加书签"""
    # 获取段落的 XML
    p = paragraph._p

    # 创建书签开始标记
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(uuid.uuid4()))
    bookmark_start.set(qn('w:name'), bookmark_name)

    # 创建书签结束标记
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))

    # 将书签标记添加到段落的 XML 中
    p.append(bookmark_start)
    p.append(bookmark_end)


def add_bookmark(paragraph, bookmark_name, position=None, additional_text=None):
    """在段落中的指定位置添加书签，并在段落末尾添加额外的文字内容"""
    # 获取段落的 XML
    p = paragraph._p

    # 创建书签开始标记
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(uuid.uuid4()))
    bookmark_start.set(qn('w:name'), bookmark_name)

    # 创建书签结束标记
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))

    # 如果指定了位置，则在该位置插入书签
    if position is not None:
        # 获取段落中的所有 run
        runs = list(p.iterchildren())
        # 计算插入位置
        insert_index = 0
        for run in runs:
            run_text = run.text or ''
            if len(run_text) > position:
                break
            position -= len(run_text)
            insert_index += 1
        p.insert(insert_index, bookmark_start)
        p.insert(insert_index + 1, bookmark_end)
    else:
        # 否则在段落末尾插入书签
        p.append(bookmark_start)
        p.append(bookmark_end)

    # 在段落末尾添加额外的文字内容
    if additional_text:
        run = paragraph.add_run(additional_text)
        run.font.color.rgb = RGBColor(0, 0, 255)


def get_header_footer(doc, section_index):
    """获取指定节的页眉和页脚"""
    section = doc.sections[section_index]
    header = section.header
    footer = section.footer
    return header, footer

def extract_page_numbers(doc):
    """从页眉中提取页码编号"""
    page_numbers = {}
    for i, section in enumerate(doc.sections):
        header, _ = get_header_footer(doc, i)
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if run.text.isdigit():
                    page_numbers[i] = int(run.text)
                    break
    return page_numbers


def mark_common_text_in_word(doc1_path, doc2_path, output1_path, output2_path, min_length=15):
    matches = compare_docs_with_threshold(doc1_path, doc2_path, min_length)
    
    # 新增处理状态追踪字典
    doc1_processed = defaultdict(set)
    doc2_processed = defaultdict(set)


     # 读取页码编号
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)
    page_numbers1 = extract_page_numbers(doc1)
    page_numbers2 = extract_page_numbers(doc2)
    


    # 处理文档1（仅处理未被处理过的段落）
    try:
        doc1 = Document(doc1_path)
        for match in matches:
            para_idx = match["doc1_para"]
            start, end = match["doc1_pos"]

            
            # 跳过已处理段落
            if para_idx in doc1_processed:
                continue
             # 确保索引有效
            if para_idx < len(doc1.paragraphs):
                if match["type"] == "full_match":
                    highlight_text(doc1.paragraphs[para_idx], 0, len(doc1.paragraphs[para_idx].text), WD_COLOR_INDEX.TURQUOISE)
                    doc1_processed[para_idx].add("full")  # 标记为已处理
                    #add_comment(doc1.paragraphs[para_idx], f"匹配于{doc2_path}文档第{match[page_numbers2]}页, 第{match['doc2_line']}行")
                    bookmark_name = f"Bk_{uuid.uuid4().hex[:4]}"  # 缩短书签名称
                    add_bookmark(doc1.paragraphs[para_idx], bookmark_name, position=end, additional_text=f"     第{match['doc2_page']}页, 第{match['doc2_line']}行")

                elif match["type"] == "partial_match" and "full" not in doc1_processed[para_idx]:
                    highlight_text(doc1.paragraphs[para_idx], start, end, WD_COLOR_INDEX.YELLOW)
                    doc1_processed[para_idx].add("partial")
                    #add_comment(doc1.paragraphs[para_idx], f"部分匹配于{doc2_path}文档第{page_numbers2}页, 第{match['doc2_line']}行")
                    bookmark_name = f"Bk_{uuid.uuid4().hex[:4]}"  # 缩短书签名称
                    add_bookmark(doc1.paragraphs[para_idx], bookmark_name, position=end, additional_text=f"     第{match['doc2_page']}页, 第{match['doc2_line']}行")
            doc1.save(output1_path)
    except Exception as e:
        print(f"❌ 处理文档1失败: {e}")

    # 处理文档2（逻辑同上）
    try:
        doc2 = Document(doc2_path)
        for match in matches:
            para_idx = match["doc2_para"]
            start, end = match["doc2_pos"]
             # 打印调试信息
            
            if para_idx in doc2_processed:
                continue
              # 确保索引有效
               
            if para_idx < len(doc2.paragraphs):
              
                if match["type"] == "full_match":
                    highlight_text(doc2.paragraphs[para_idx], 0, len(doc2.paragraphs[para_idx].text), WD_COLOR_INDEX.TURQUOISE)
                    doc2_processed[para_idx].add("full")
                    add_comment(doc2.paragraphs[para_idx], f"完全匹配于文档1第{match['doc1_page']}页, 第{match['doc1_line']}行")
                    bookmark_name = f"Bk_{uuid.uuid4().hex[:4]}"  # 缩短书签名称
                    add_bookmark(doc2.paragraphs[para_idx], bookmark_name, position=end, additional_text=f"    第{match['doc1_page']}页, 第{match['doc1_line']}行")
                elif match["type"] == "partial_match" and "full" not in doc2_processed[para_idx]:
                    highlight_text(doc2.paragraphs[para_idx], start, end, WD_COLOR_INDEX.YELLOW)
                    doc2_processed[para_idx].add("partial")
                    add_comment(doc2.paragraphs[para_idx], f"部分匹配于文档1第{match['doc1_page']}页, 第{match['doc1_line']}行")
                    bookmark_name = f"Bk_{uuid.uuid4().hex[:4]}"  # 缩短书签名称
                    add_bookmark(doc2.paragraphs[para_idx], bookmark_name, position=end, additional_text=f"    第{match['doc1_page']}页, 第{match['doc1_line']}行")
            doc2.save(output2_path)
    except Exception as e:
        print(f"❌ 处理文档2失败: {e}")

# ------------------ 日志配置 ------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("PDFComparator")

# ================== PDF对比模块 ==================

def compare_pdfs(file1, file2, output_dir, min_length):
    try:
        # 检查输入文件是否存在
        if not os.path.exists(file1):
            raise FileNotFoundError(f"文件不存在: {file1}")
        if not os.path.exists(file2):
            raise FileNotFoundError(f"文件不存在: {file2}")

        
        # 定义基础路径
        base_path = "D:/Compare/static"
        
        # 假设 file1 和 file2 是你已经定义的文件路径
        rel_file1 = os.path.relpath(file1, base_path)
        rel_file2 = os.path.relpath(file2, base_path)
        
        # 在相对路径前加上 ./ 
        rel_file1 = f"{rel_file1}"
        rel_file2 = f"{rel_file2}"
        
        # 将相对路径写入 static/paths.txt 文件
        with open('static/paths.txt', 'w') as f:
            f.write(f"{rel_file1}\n{rel_file2}")

        # 检查输出目录是否存在，不存在则创建
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_file = os.path.join(output_dir, "CommonParagraphs.json")

        def extract_paragraphs(file_path, min_length):
            paragraphs = []
            with pdfplumber.open(file_path) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')  # 按行分割文本
                        for line_number, line in enumerate(lines, start=1):
                            if len(line) >= min_length:
                                paragraphs.append({
                                    'page': page_number,
                                    'line': line_number,
                                    'text': line
                                })
            return paragraphs


        # 提取段落
        paragraphs1 = extract_paragraphs(file1, min_length)
        paragraphs2 = extract_paragraphs(file2, min_length)

        # 比较段落
        common_paragraphs = []
        seen_common_substrings = set()
            
        for para1 in paragraphs1:
            for para2 in paragraphs2:
                # 去除特殊字符
                clean_text1 = remove_special_chars(para1['text'])
                clean_text2 = remove_special_chars(para2['text'])
        
                common_substrings = find_common_substrings(clean_text1, clean_text2, min_length)
                if common_substrings:
                    for substring in common_substrings:
                        if (para1['page'], para1['line'], para2['page'], para2['line'], substring) not in seen_common_substrings:
                            seen_common_substrings.add((para1['page'], para1['line'], para2['page'], para2['line'], substring))
                            common_paragraphs.append({
                                'file1': file1,
                                'page1': para1['page'],
                                'line1': para1['line'],
                                'text1': para1['text'],
                                'file2': file2,
                                'page2': para2['page'],
                                'line2': para2['line'],
                                'text2': para2['text'],
                                'common_substrings': [substring]
                            })


        # 将结果保存到 JSON 文件
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'paragraphs1': paragraphs1,
                'paragraphs2': paragraphs2,
                'common_paragraphs': common_paragraphs
            }, f, ensure_ascii=False, indent=4)

    except Exception as e:
        print(f"发生错误: {e}")


def find_common_substrings(str1, str2, min_length):
    matcher = SequenceMatcher(None, str1, str2)
    common_substrings = []
    for match in matcher.get_matching_blocks():
        if match.size >= min_length:
            common_substrings.append(str1[match.a:match.a + match.size])
    return common_substrings

def remove_special_chars(text):
    # 保留中文字符，去除其他特殊字符
    return re.sub(r'[^\u4e00-\u9fff]', '', text)    
    
# 图像对比模块  
def extract_images_from_pdf(pdf_path, output_folder):
    # 打开 PDF 文件
    pdf_document = fitz.open(pdf_path)
    
    # 确保输出目录存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 遍历每一页
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        
        # 获取页面上的所有图像
        image_list = page.get_images(full=True)
        
        # 遍历每个图像
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            # 构建图像文件名
            image_filename = f"page_{page_num + 1}_img_{img_index + 1}.{image_ext}"
            image_path = os.path.join(output_folder, image_filename)
            
            # 保存图像
            with open(image_path, "wb") as image_file:
                image_file.write(image_bytes)
            
            print(f"Extracted image: {image_path}")

    
# ================== 主控制流程 ==================
def process_files(file1: str, file2: str, output_dir: str, min_length: int):
    """统一处理入口"""
    def get_ext(path: str) -> str:
        return os.path.splitext(path)[1].lower()
    
    if get_ext(file1) != get_ext(file2):
        raise ValueError("文件格式不匹配")
    
    file_type = get_ext(file1)
    base1 = os.path.basename(file1).split('.')[0]
    base2 = os.path.basename(file2).split('.')[0]
    
    if file_type == '.docx':
        output1 = os.path.join(output_dir, f"{base1}_compared.docx")
        output2 = os.path.join(output_dir, f"{base2}_compared.docx")
        mark_common_text_in_word(file1, file2, output1, output2, min_length)
    elif file_type == '.pdf':
        output = os.path.join(output_dir, "JsonFromPdf")
        compare_pdfs(file1, file2, output, min_length)
        
    else:
        raise ValueError(f"不支持的格式: {file_type}")
# min_length 内容对比阈值  output 输出路径
def main():
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口

    # 选择第一个 PDF 文件
    file1 = filedialog.askopenfilename(title="选择第一个 PDF 文件", filetypes=[("PDF files", "*.pdf")])
    if not file1:
        messagebox.showerror("错误", "未选择第一个 PDF 文件")
        return

    # 选择第二个 PDF 文件
    file2 = filedialog.askopenfilename(title="选择第二个 PDF 文件", filetypes=[("PDF files", "*.pdf")])
    if not file2:
        messagebox.showerror("错误", "未选择第二个 PDF 文件")
        return

    # 选择输出目录
    output_dir = filedialog.askdirectory(title="选择输出目录")
    if not output_dir:
        messagebox.showerror("错误", "未选择输出目录")
        return

    # 获取最小匹配长度
    min_length = tk.simpledialog.askinteger("输入", "请输入最小匹配长度（默认：13字符 ）", initialvalue=13)
    if min_length is None:
        min_length = 13

    try:
        process_files(file1, file2, output_dir, min_length)
        messagebox.showinfo("成功", f"处理完成！结果保存在：{output_dir}")
    except Exception as e:
        messagebox.showerror("错误", f"处理失败：{e}")

if __name__ == "__main__":
    main()